#!/usr/bin/env python3
"""Balise le livret CIL officiel 2024 (.docx) avec les tags docxtemplater.

Produit `public/cil/template.docx` à partir du .docx officiel. Reproductible :
relancer si le modèle officiel évolue. Conventions (comme le RCI) :
  {txt_x}         → texte
  {check_x}       → ☒ / ☐ (piloté par la donnée)
  {num_NN}        → numéro 10-69 (barré si utilisé)
  {%photo_sig_x}  → image de signature (base64)

Usage : python scripts/cil-tag-template.py <source.docx> [public/cil/template.docx]

Périmètre balisé (Lot 2) : en-tête, établissement, réglettes de numéros,
sous-tables N° donné/reçu des protections (CRC/RSS/AC) et reprises/
rétablissements, horaires d'arrivée/départ des intervenants, changement de CIL,
signatures. Différé : blancs inline des phrases de dépêche + détail du carnet
libre (voir plan).
"""
import re
import sys

import docx
from docx.oxml.ns import qn


def set_cell(cell, text):
    """Remplace le contenu d'une cellule par un unique run (tag contigu)."""
    cell.text = text  # setter python-docx : vide la cellule + un paragraphe/run


def tag_numbers(document):
    """Toute cellule dont le texte est un entier 10-69 → {num_NN}."""
    for table in _all_tables(document):
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if re.fullmatch(r"\d{2}", t) and 10 <= int(t) <= 69:
                    set_cell(cell, f"{{num_{t}}}")


def _all_tables(document):
    """Tables de 1er niveau + tables imbriquées (récursif)."""
    out = []
    stack = list(document.tables)
    while stack:
        t = stack.pop()
        out.append(t)
        for row in t.rows:
            for cell in row.cells:
                stack.extend(cell.tables)
    return out


def nested(cell):
    return cell.tables[0] if cell.tables else None


def set_para(p, text):
    """Réécrit un paragraphe en un unique run (tag contigu), en conservant la
    police du 1ᵉʳ run existant (le gras décoratif de la phrase est préservé)."""
    runs = list(p.runs)
    for r in runs[1:]:
        r.text = ""
    if runs:
        runs[0].text = text
    else:
        p.add_run(text)


def _mettre_a_echelle_table(table, facteur):
    """Applique un facteur aux largeurs d'une table ET de ses tables imbriquées.

    La récursion est indispensable : une sous-table laissée à sa taille d'origine
    élargit sa cellule parente, et toute la table déborde malgré le
    redimensionnement de la grille.
    """
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for c in grid.findall(qn("w:gridCol")):
            c.set(qn("w:w"), str(max(1, round(int(c.get(qn("w:w"))) * facteur))))

    tblW = table._tbl.tblPr.find(qn("w:tblW"))
    if tblW is not None and tblW.get(qn("w:type")) == "dxa":
        tblW.set(qn("w:w"), str(max(1, round(int(tblW.get(qn("w:w"))) * facteur))))

    vues = set()
    for row in table.rows:
        for cell in row.cells:
            if cell._tc in vues:
                continue
            vues.add(cell._tc)
            # Les cellules portent leur propre largeur : Word la privilégie,
            # redimensionner la seule grille resterait sans effet.
            tcPr = cell._tc.tcPr
            tcW = tcPr.find(qn("w:tcW")) if tcPr is not None else None
            if tcW is not None and tcW.get(qn("w:type")) == "dxa":
                tcW.set(
                    qn("w:w"), str(max(1, round(int(tcW.get(qn("w:w"))) * facteur)))
                )
            for sous_table in cell.tables:
                _mettre_a_echelle_table(sous_table, facteur)


def largeur_table_mm(table):
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is None:
        return 0.0
    return sum(int(c.get(qn("w:w"))) for c in grid.findall(qn("w:gridCol"))) / 1440 * 25.4


def ajuster_largeur_table(table, section):
    """Ramène une table à la largeur imprimable en gardant ses proportions.

    Un facteur unique est appliqué à la grille, aux cellules et aux tables
    imbriquées : les rapports de largeur du modèle officiel sont préservés,
    seule la largeur totale change. Renvoie la largeur obtenue, en mm.
    """
    actuelle_mm = largeur_table_mm(table)
    if actuelle_mm == 0:
        return None
    utile_emu = section.page_width - section.left_margin - section.right_margin
    cible_mm = utile_emu / 635 / 1440 * 25.4
    if actuelle_mm <= cible_mm:
        return actuelle_mm
    _mettre_a_echelle_table(table, cible_mm / actuelle_mm)
    return largeur_table_mm(table)


def passer_en_a4_portrait(document, facteur=0.60):
    """Repagine le livret en A4 portrait : 2 pages (livret, puis carnet).

    C'est l'agencement du livret papier d'origine. Contrairement à l'A3 — qui
    partage sa largeur avec l'A4 paysage — il faut ici réduire AUSSI les
    tableaux : la largeur utile tombe de 271,6 à 184,6 mm.

    Le facteur s'applique aux polices et aux hauteurs de ligne ; les tableaux
    sont ensuite ramenés à la largeur utile (cf. `ajuster_largeur_table`), ce
    qui préserve leurs proportions internes. Les deux réductions doivent rester
    du même ordre : une police trop grande dans un tableau rétréci provoquerait
    des retours à la ligne, et donc un débordement en hauteur.

    Le facteur 0,60 (12 pt → 7,2 pt) est le plus grand qui tienne en 2 pages
    avec ~20 mm de réserve : au-delà, un texte de dépêche un peu long fait
    basculer sur une 3ᵉ page. C'est la densité de l'imprimé papier d'origine.

    Le document conservant deux sections (livret, puis carnet), le résultat
    tient en 2 pages.
    """
    from docx.enum.section import WD_ORIENT
    from docx.shared import Mm, Pt

    for section in document.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(210)
        section.page_height = Mm(297)

    # Style de base (les runs sans taille explicite en héritent).
    normal = document.styles["Normal"]
    if normal.font.size is not None:
        normal.font.size = Pt(round(normal.font.size.pt * facteur, 1))

    runs = 0
    for rPr in document.element.body.iter(qn("w:rPr")):
        sz = rPr.find(qn("w:sz"))
        if sz is not None:
            sz.set(qn("w:val"), str(max(2, round(int(sz.get(qn("w:val"))) * facteur))))
            runs += 1
        szCs = rPr.find(qn("w:szCs"))
        if szCs is not None:
            szCs.set(
                qn("w:val"), str(max(2, round(int(szCs.get(qn("w:val"))) * facteur)))
            )

    # Hauteurs de ligne : ce sont des minimums, ils empêcheraient la réduction.
    lignes = 0
    for trPr in document.element.body.iter(qn("w:trPr")):
        h = trPr.find(qn("w:trHeight"))
        if h is not None:
            h.set(qn("w:val"), str(max(1, round(int(h.get(qn("w:val"))) * facteur))))
            lignes += 1

    # Tous les tableaux de 1er niveau sont ramenés dans la largeur utile.
    section = document.sections[0]
    tables = 0
    for table in document.tables:
        avant = table._tbl.find(qn("w:tblGrid"))
        if avant is None:
            continue
        largeur = sum(int(c.get(qn("w:w"))) for c in avant.findall(qn("w:gridCol")))
        if largeur / 1440 * 25.4 > 184.7:
            ajuster_largeur_table(table, section)
            tables += 1
    return runs, lignes, tables


def tag_sentences(cell, rules):
    """Applique à chaque paragraphe la 1ʳᵉ règle (predicat, texte) qui matche.

    `texte` peut être une chaîne ou un callable recevant le texte d'origine,
    pour les phrases dont la coupure varie d'un cadre à l'autre.
    """
    for p in cell.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        for pred, new_text in rules:
            if pred(t):
                set_para(p, new_text(t) if callable(new_text) else new_text)
                break


def main():
    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else "public/cil/template.docx"
    d = docx.Document(src)
    T = d.tables

    # ── En-tête (table 0) ────────────────────────────────────────────────────
    set_cell(T[0].rows[1].cells[1], "{txt_dateHeure}")
    set_cell(T[0].rows[2].cells[1], "{txt_incident}")
    set_cell(T[0].rows[3].cells[1], "{txt_lieu}")
    set_cell(T[0].rows[4].cells[1], "{txt_voie}")
    set_cell(T[0].rows[1].cells[4], "{txt_nomCil}")
    set_cell(
        T[0].rows[2].cells[4],
        "{check_etab_ral} EIC RAL    {check_etab_rhn} INFP RHN    {check_etab_lgv} INFP LGV",
    )
    set_cell(T[0].rows[3].cells[4], "{txt_designeA}")

    # ── Protections (table 1) : sous-tables N° donné/reçu/Date-Heure ─────────
    elec = nested(T[1].rows[1].cells[0])
    set_cell(elec.rows[1].cells[1], "{txt_prot_elec_crc_donne}")
    set_cell(elec.rows[1].cells[2], "{txt_prot_elec_crc_recu}")
    set_cell(elec.rows[1].cells[3], "{txt_prot_elec_crc_dh}")
    set_cell(elec.rows[2].cells[1], "{txt_prot_elec_rss_donne}")
    set_cell(elec.rows[2].cells[2], "{txt_prot_elec_rss_recu}")
    set_cell(elec.rows[2].cells[3], "{txt_prot_elec_rss_dh}")
    circ = nested(T[1].rows[1].cells[4])
    set_cell(circ.rows[1].cells[1], "{txt_prot_circ_crc_donne}")
    set_cell(circ.rows[1].cells[2], "{txt_prot_circ_crc_recu}")
    set_cell(circ.rows[1].cells[3], "{txt_prot_circ_crc_dh}")
    # Libellé de la 2ᵉ ligne : « AC de … » (destinataire réel de la dépêche).
    set_cell(circ.rows[2].cells[0], "AC de {txt_ac_label}")
    set_cell(circ.rows[2].cells[1], "{txt_prot_circ_ac_donne}")
    set_cell(circ.rows[2].cells[2], "{txt_prot_circ_ac_recu}")
    set_cell(circ.rows[2].cells[3], "{txt_prot_circ_ac_dh}")
    # Avis COS / OPJ (dates)
    set_cell(T[1].rows[2].cells[2], "{txt_prot_elec_avis_cos}")
    set_cell(T[1].rows[3].cells[2], "{txt_prot_elec_avis_opj}")
    set_cell(T[1].rows[2].cells[6], "{txt_prot_circ_avis_cos}")
    set_cell(T[1].rows[3].cells[6], "{txt_prot_circ_avis_opj}")

    # ── Réglette 10-29 + Changement de CIL + intervenants arrivée (table 2) ──
    set_cell(T[2].rows[4].cells[16], "{txt_chg_ac}")
    set_cell(T[2].rows[4].cells[21], "{txt_chg_crc}")
    set_cell(T[2].rows[4].cells[25], "{txt_chg_opj}")
    set_cell(T[2].rows[4].cells[29], "{txt_chg_cos}")
    arr = T[2].rows[8]
    for col, key in (
        (2, "cos"), (5, "opj"), (8, "pf"), (11, "eic"),
        (14, "infp"), (18, "exft"), (22, "exfv"),
    ):
        set_cell(arr.cells[col], f"{{txt_arr_{key}}}")

    # ── Départ effectif des intervenants (table 4) ──────────────────────────
    dep = T[4].rows[1]
    for col, key in (
        (1, "cos"), (2, "opj"), (3, "pf"), (4, "eic"),
        (5, "infp"), (6, "exft"), (7, "exfv"),
    ):
        set_cell(dep.cells[col], f"{{txt_dep_{key}}}")

    # ── Reprises / rétablissements (tables 3 partiel, 5 normal) ─────────────
    def tag_reprise(table, retab_prefix, reprise_prefix):
        rn = nested(table.rows[4].cells[0])  # rétab (gauche)
        set_cell(rn.rows[1].cells[0], f"{{txt_{retab_prefix}_donne}}")
        set_cell(rn.rows[1].cells[1], f"{{txt_{retab_prefix}_recu}}")
        set_cell(rn.rows[1].cells[2], f"{{txt_{retab_prefix}_dh}}")
        rp = nested(table.rows[4].cells[5])  # reprise (droite)
        set_cell(rp.rows[1].cells[0], f"{{txt_{reprise_prefix}_donne}}")
        set_cell(rp.rows[1].cells[1], f"{{txt_{reprise_prefix}_recu}}")
        set_cell(rp.rows[1].cells[2], f"{{txt_{reprise_prefix}_dh}}")
        # Chaque autorité signe SUR SA PROPRE LIGNE d'autorisation : la cellule
        # immédiatement à droite de « Autorisation reçue du COS/OPJ : le … ».
        # (« Signature » en haut n'est que l'en-tête de colonne.)
        #
        # ATTENTION : dans cette grille fusionnée irrégulière, python-docx renvoie
        # le MÊME objet cellule pour des (ligne, colonne) différents — on ne peut
        # pas cibler par indices. On parcourt donc les cellules DISTINCTES de
        # chaque ligne et on prend celle qui suit le libellé d'autorisation.
        for row in table.rows:
            distinctes, vues = [], set()
            for c in row.cells:
                if c._tc in vues:
                    continue
                vues.add(c._tc)
                distinctes.append(c)
            # Cadre gauche = rétablissement, cadre droit = reprise : les libellés
            # apparaissent dans cet ordre sur la ligne.
            role = None
            for i, c in enumerate(distinctes):
                t = c.text.strip()
                if t.startswith("Autorisation reçue du COS"):
                    role = "cos"
                elif t.startswith("Autorisation reçue du OPJ"):
                    role = "opj"
                else:
                    continue
                if i + 1 >= len(distinctes):
                    continue
                # 1ʳᵉ occurrence = cadre gauche (rétablissement), 2ᵉ = droite.
                pfx = retab_prefix if i == 0 else reprise_prefix
                set_cell(distinctes[i + 1], "{%photo_sig_" + pfx + "_" + role + "}")

    tag_reprise(T[3], "retp", "repp")  # partiels
    tag_reprise(T[5], "retn", "repn")  # normaux

    # ── Blancs inline des PHRASES de dépêche (fix « rien ne remonte ») ───────
    sw = lambda s: (lambda t: t.startswith(s))  # noqa: E731
    contains = lambda s: (lambda t: s in t)  # noqa: E731

    tag_sentences(T[1].rows[1].cells[0], [  # protection électrique (phrase CRC)
        (sw("M."), "M. {txt_prot_elec_cil} CIL à CRC de Lyon :"),
        (sw("Je reprends"), "Je reprends à mon compte la coupure d'urgence demandée suite à (l'événement) {txt_prot_elec_evenement} sur voie(s) {txt_prot_elec_voies} au kilomètre {txt_prot_elec_km}"),
        (sw("entre les gares"), "{check_gare_entre} entre les gares de {txt_gareA} et de {txt_gareB}"),
        (sw("ou"), "ou {check_gare_engare} en gare de {txt_gareUnique}"),
        (sw("Motif"), "Motif : {txt_prot_elec_motif}"),
    ])
    tag_sentences(T[1].rows[1].cells[4], [  # protection circulation
        (sw("M."), "M. {txt_prot_circ_cil} CIL à CRC de Lyon :"),
        (sw("Je reprends"), "Je reprends à mon compte les mesures de protection suite à (l'événement) {txt_prot_circ_evenement} sur voie(s) {txt_prot_circ_voies} au km {txt_prot_circ_km}"),
        (sw("entre les gares"), "{check_gare_entre} entre les gares de {txt_gareA} et de {txt_gareB}"),
        (sw("ou"), "ou {check_gare_engare} en gare de {txt_gareUnique}"),
        (contains("interdites"), "Voies n° {txt_prot_circ_vi} interdites à la circulation"),
        (contains("prudente"), "Voies n° {txt_prot_circ_vp} circulation en marche prudente"),
        (contains("normale"), "Voies n° {txt_prot_circ_vn} circulation en marche normale"),
        (sw("Motif"), "Motif : {txt_prot_circ_motif}"),
    ])

    def sentence_reprise(cell, pfx):
        # ATTENTION : la phrase des cadres PARTIELS tient sur une ligne
        # (« sur voie(s) … aux abords du kilomètre … situé ») alors que celle des
        # cadres NORMAUX est coupée : « …sur voie(s) » termine une ligne et les
        # blancs sont sur la suivante, qui commence par des pointillés. On
        # matche donc sur « aux abords du kilom » et on ne réécrit « sur voie(s) »
        # que si la ligne le portait déjà.
        def ligne_km(t):
            prefixe = "sur voie(s) " if t.startswith("sur voie") else ""
            return (
                prefixe
                + "{txt_%s_voies} aux abords du kilomètre {txt_%s_km} situé" % (pfx, pfx)
            )

        tag_sentences(cell, [
            (sw("M."), "M. {txt_%s_cil} CIL à AC de {txt_%s_ac}" % (pfx, pfx)),
            (sw("Au titre"), "Au titre de (l'événement) {txt_%s_evenement}" % pfx),
            (contains("aux abords du kilom"), ligne_km),
            (sw("entre les gares"), "{check_gare_entre} entre les gares de {txt_gareA} et de {txt_gareB}"),
            (sw("ou"), "ou {check_gare_engare} en gare de {txt_gareUnique}"),
        ])

    def sentence_retab(cell, pfx):
        tag_sentences(cell, [
            (sw("M."), "M. {txt_%s_cil} CIL à RSS de Lyon :" % pfx),
            (sw("sur voie"), "sur voie(s) {txt_%s_voies} au kilomètre {txt_%s_km}" % (pfx, pfx)),
        ])

    sentence_reprise(T[3].rows[4].cells[5], "repp")
    sentence_reprise(T[5].rows[4].cells[5], "repn")
    sentence_retab(T[3].rows[4].cells[0], "retp")
    sentence_retab(T[5].rows[4].cells[0], "retn")

    # ── Autorisations COS/OPJ + Avis au CRC (reprises / rétablissements) ────
    def tag_autorisations(table, retab_pfx, reprise_pfx):
        for col, pfx in ((0, retab_pfx), (5, reprise_pfx)):
            set_para(
                table.rows[2].cells[col].paragraphs[0],
                "Autorisation reçue du COS : {txt_%s_autor_cos}" % pfx,
            )
            set_para(
                table.rows[3].cells[col].paragraphs[0],
                "Autorisation reçue du OPJ : {txt_%s_autor_opj}" % pfx,
            )
        # Avis au CRC : la phrase pré-imprimée se termine par « à ……… », qui
        # n'attend QUE l'heure (contrairement aux avis COS/OPJ des protections,
        # dont la case porte « le ../../.. à ..h.. » et veut date + heure).
        # On remplit donc le blanc de la phrase, et on vide la colonne de droite
        # pour ne pas afficher deux fois la même information.
        cellules = []
        for row in table.rows:
            vues = set()
            for c in row.cells:
                if c._tc in vues:
                    continue
                vues.add(c._tc)
                if c.text.strip().startswith("au CRC de Lyon"):
                    cellules.append((row, c))
        assert len(cellules) == 2, f"cellules « Avis au CRC » attendues: 2, trouvées {len(cellules)}"
        # 1ʳᵉ occurrence = cadre gauche (rétablissement), 2ᵉ = droite (reprise).
        for (row, cell), pfx in zip(cellules, (retab_pfx, reprise_pfx)):
            paras = cell.paragraphs
            # La 2ᵉ ligne porte « à ……… » : c'est elle qui reçoit l'heure.
            if len(paras) > 1:
                set_para(paras[1], "à {txt_%s_avis_crc}" % pfx)
            else:
                cell.add_paragraph("à {txt_%s_avis_crc}" % pfx)
            # Colonne de droite (« le ../../.. à ..h.. ») : vidée.
            vues, suivante = set(), None
            apres = False
            for c in row.cells:
                if c._tc in vues:
                    continue
                vues.add(c._tc)
                if apres:
                    suivante = c
                    break
                if c._tc is cell._tc:
                    apres = True
            if suivante is not None:
                for para in suivante.paragraphs:
                    set_para(para, "")

    tag_autorisations(T[3], "retp", "repp")
    tag_autorisations(T[5], "retn", "repn")

    # ── N° de téléphone des intervenants (ligne « N° de tél ») ──────────────
    tel_row = T[2].rows[9]
    for col, key in (
        (2, "cos"), (5, "opj"), (8, "pf"), (11, "eic"),
        (14, "infp"), (18, "exft"), (22, "exfv"),
    ):
        set_cell(tel_row.cells[col], f"{{txt_tel_{key}}}")

    # ── Changement de CIL (phrase) ─────────────────────────────────────────
    set_para(
        T[2].rows[4].cells[0].paragraphs[0],
        "M. {txt_chg_cil} CIL, je suis remplacé par M. {txt_chg_remplacant} à {txt_chg_heure}",
    )

    # ── Carnet d'enregistrement des dépêches (lignes libres, non numérotées
    #    à l'impression : le n° est écrit par le CIL) ──────────────────────
    carnet = T[7]
    COLS = {"num": 0, "date": 2, "de": 4, "texte": 7, "expedie": 20, "recu": 23, "heure": 25}
    line = 0
    for ri in range(1, len(carnet.rows)):
        row = carnet.rows[ri]
        # La dernière ligne est la réglette de numéros (50-69) : on l'ignore.
        if any(c.text.strip().isdigit() for c in row.cells):
            continue
        line += 1
        for key, col in COLS.items():
            set_cell(row.cells[col], f"{{txt_libre_{line}_{key}}}")
    print("lignes de carnet balisées :", line)


    # ── Réglettes de numéros (générique, en dernier) ────────────────────────
    tag_numbers(d)

    # Repagination en dernier : elle fixe le format, l'échelle du texte et la
    # largeur des tableaux d'un seul tenant.
    runs, lignes, tables = passer_en_a4_portrait(d)
    print(
        "A4 portrait : %d runs, %d hauteurs de ligne, %d tableaux ajustés"
        % (runs, lignes, tables)
    )

    d.save(out)
    print("template écrit :", out)


if __name__ == "__main__":
    main()
